"""
Aplikasi Pembuatan Surat Penagihan - SRR Kalibata
Mengambil data dari Google Sheets / upload file dan mengisi template DOCX
Versi: 1.1 (Streamlit Cloud-ready)
"""

import io
import re
import zipfile
from datetime import date, datetime

import pandas as pd
import requests
import streamlit as st
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ─── PAGE CONFIG ─────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Surat Penagihan SRR",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─── CONSTANTS ────────────────────────────────────────────────────────────────
SPREADSHEET_ID = "1imGyzAXvznWgw5Ph8EyyIuFPEdZziv7TBJqjXwrzZ5M"

BANK_OPTIONS = {
    "Bank Mandiri — 126-0005748719": {
        "bank": "Bank Mandiri KCP JKT Kalibata Rawajati",
        "norek": "126-0005748719",
    },
    "Bank JTrust — 1001883933": {
        "bank": "Bank JTrust Indonesia",
        "norek": "1001883933",
    },
    "BRI — 042601000618306": {
        "bank": "Bank Rakyat Indonesia Jkt Kalibata",
        "norek": "042601000618306.",
    },
    "BNI — 0981981462": {
        "bank": "BNI, Kantor Cabang Tebet",
        "norek": "0981981462.",
    },
}

BULAN_ID = {
    1: "Januari", 2: "Februari", 3: "Maret", 4: "April",
    5: "Mei", 6: "Juni", 7: "Juli", 8: "Agustus",
    9: "September", 10: "Oktober", 11: "November", 12: "Desember",
}

# ─── CSS ─────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
.main-header {
    background: linear-gradient(135deg, #1a5276 0%, #2e86c1 100%);
    color: white; padding: 1.2rem 1.5rem; border-radius: 10px;
    margin-bottom: 1.5rem; box-shadow: 0 4px 15px rgba(0,0,0,0.15);
}
.main-header h1 { margin: 0; font-size: 1.6rem; }
.main-header p  { margin: 0.3rem 0 0; opacity: 0.85; font-size: 0.9rem; }
.section-card {
    background: #f8f9fa; border-left: 4px solid #2e86c1;
    padding: 0.9rem 1.1rem; border-radius: 0 8px 8px 0; margin-bottom: 1rem;
}
.section-card h3 { margin: 0 0 0.7rem; color: #1a5276; font-size: 0.95rem; }
.preview-box {
    background: #f0f7fd; border: 1px solid #aed6f1;
    border-radius: 8px; padding: 0.9rem 1rem; font-size: 0.88rem;
}
.saved-badge {
    display:inline-block; background:#2e86c1; color:white;
    padding:2px 8px; border-radius:12px; font-size:0.72rem; margin-left:6px;
}
</style>
""", unsafe_allow_html=True)

# ─── HELPER: TERBILANG ────────────────────────────────────────────────────────
def terbilang(n: int) -> str:
    if n < 0:
        return "minus " + terbilang(-n)
    if n == 0:
        return "nol"
    satuan = [
        "", "satu", "dua", "tiga", "empat", "lima", "enam", "tujuh", "delapan",
        "sembilan", "sepuluh", "sebelas", "dua belas", "tiga belas", "empat belas",
        "lima belas", "enam belas", "tujuh belas", "delapan belas", "sembilan belas",
    ]

    def _b(x):
        if x == 0:
            return ""
        if x < 20:
            return satuan[x]
        if x < 100:
            tens = x // 10
            ones = x % 10
            prefix = (
                "dua puluh tiga puluh empat puluh lima puluh "
                "enam puluh tujuh puluh delapan puluh sembilan puluh"
            ).split()
            t = prefix[(tens - 2) * 2 : (tens - 2) * 2 + 2]
            return " ".join(t) + (" " + satuan[ones] if ones else "")
        if x < 200:
            rest = x - 100
            return "seratus" + (" " + _b(rest) if rest else "")
        if x < 1_000:
            h = x // 100
            rest = x % 100
            return satuan[h] + " ratus" + (" " + _b(rest) if rest else "")
        if x < 2_000:
            rest = x - 1_000
            return "seribu" + (" " + _b(rest) if rest else "")
        if x < 1_000_000:
            k = x // 1_000
            rest = x % 1_000
            return _b(k) + " ribu" + (" " + _b(rest) if rest else "")
        if x < 1_000_000_000:
            m = x // 1_000_000
            rest = x % 1_000_000
            return _b(m) + " juta" + (" " + _b(rest) if rest else "")
        if x < 1_000_000_000_000:
            b = x // 1_000_000_000
            rest = x % 1_000_000_000
            return _b(b) + " miliar" + (" " + _b(rest) if rest else "")
        t2 = x // 1_000_000_000_000
        rest = x % 1_000_000_000_000
        return _b(t2) + " triliun" + (" " + _b(rest) if rest else "")

    w = _b(n).strip()
    return w[0].upper() + w[1:] if w else w


def terbilang_rupiah(amount: float) -> str:
    """Terbilang dengan Title Case setiap kata."""
    w = terbilang(round(amount))
    w_titled = " ".join(word.capitalize() for word in w.split())
    return f"{w_titled} Rupiah"


def format_rp(amount: float) -> str:
    return f"{int(round(amount)):,}".replace(",", ".")


def format_tgl(d: date) -> str:
    """Tanggal surat header — hanya tanggal, tanpa prefix 'Jakarta, '."""
    return f"{d.day} {BULAN_ID[d.month]} {d.year}"


def extract_kode_pt(no_proposal: str) -> str:
    """Ambil Kode_PT dari no_proposal: segmen sebelum segmen terakhir.

    Contoh: 260101.003/SRR-JK/SPN-ABF/BJTI/OR -> BJTI
    """
    parts = [p.strip() for p in str(no_proposal).split("/") if p.strip()]
    if len(parts) >= 2:
        return parts[-2]
    return ""


# ─── HELPER: LOAD DATA ────────────────────────────────────────────────────────
def load_from_csv_url(url: str) -> pd.DataFrame:
    """Muat data dari URL Google Sheets (public) sebagai CSV."""
    export_url = url.strip()
    m = re.search(r"spreadsheets/d/([A-Za-z0-9_-]+)", export_url)
    if m:
        sheet_id = m.group(1)
        # Coba ambil parameter gid jika ada (untuk sheet tertentu)
        gid_match = re.search(r"[#&?]gid=(\d+)", export_url)
        gid_param = f"&gid={gid_match.group(1)}" if gid_match else ""
        export_url = (
            f"https://docs.google.com/spreadsheets/d/{sheet_id}"
            f"/export?format=csv{gid_param}"
        )

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/120.0.0.0 Safari/537.36"
        )
    }
    resp = requests.get(export_url, headers=headers, timeout=30)
    resp.raise_for_status()

    # Deteksi jika Google mengembalikan HTML error (sheet tidak public)
    content_type = resp.headers.get("Content-Type", "")
    if "text/html" in content_type and len(resp.content) < 5000:
        raise ValueError(
            "Google Sheets mengembalikan halaman HTML, bukan CSV. "
            "Pastikan sheet sudah di-share sebagai 'Anyone with the link can view'."
        )

    return pd.read_csv(io.StringIO(resp.text))


def normalize_df(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df.columns = [c.strip().lower().replace(" ", "_") for c in df.columns]
    if "pemberi_tugas" not in df.columns:
        raise ValueError(
            f"Kolom 'pemberi_tugas' tidak ditemukan. "
            f"Kolom yang tersedia: {list(df.columns)}"
        )
    df = df[df["pemberi_tugas"].notna() & (df["pemberi_tugas"].str.strip() != "")]
    df = df.sort_values(
        "pemberi_tugas", key=lambda x: x.str.strip().str.upper()
    ).reset_index(drop=True)

    if "proposed_fee" in df.columns:
        df["proposed_fee"] = pd.to_numeric(
            df["proposed_fee"]
            .astype(str)
            .str.replace(r"[^\d.]", "", regex=True),
            errors="coerce",
        ).fillna(0)
    else:
        df["proposed_fee"] = 0.0
    return df


# ─── HELPER: DOCX FILL ────────────────────────────────────────────────────────
def _make_run_elem(text, bold, underline, italic, font_size, font_name):
    """Buat elemen w:r baru dengan teks dan formatting yang diberikan."""
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    if bold is True:
        rPr.append(OxmlElement("w:b"))
    elif bold is False:
        e = OxmlElement("w:b")
        e.set(qn("w:val"), "0")
        rPr.append(e)

    if underline is True:
        e = OxmlElement("w:u")
        e.set(qn("w:val"), "single")
        rPr.append(e)

    if italic is True:
        rPr.append(OxmlElement("w:i"))

    if font_size is not None:
        try:
            half_pts = str(int(font_size.pt * 2))
            e = OxmlElement("w:sz")
            e.set(qn("w:val"), half_pts)
            rPr.append(e)
            e = OxmlElement("w:szCs")
            e.set(qn("w:val"), half_pts)
            rPr.append(e)
        except Exception:
            pass

    if font_name:
        e = OxmlElement("w:rFonts")
        e.set(qn("w:ascii"), font_name)
        e.set(qn("w:hAnsi"), font_name)
        rPr.append(e)

    if len(rPr):
        r.append(rPr)

    # Mempertahankan Format Tab sebagai tag <w:tab/> yang sah
    if "\t" in text:
        parts = text.split("\t")
        for i, part in enumerate(parts):
            if part:
                t = OxmlElement("w:t")
                t.text = part
                if part.startswith(" ") or part.endswith(" "):
                    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                r.append(t)
            if i < len(parts) - 1:
                r.append(OxmlElement("w:tab"))
    else:
        t = OxmlElement("w:t")
        t.text = text
        if text and (text.startswith(" ") or text.endswith(" ")):
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t)

    return r


def replace_in_para(para, reps: dict):
    full = "".join(r.text for r in para.runs)
    if not any(k in full for k in reps):
        return

    char_fmts = []
    for run in para.runs:
        fmt = (run.bold, run.underline, run.italic, run.font.size, run.font.name)
        for _ in run.text:
            char_fmts.append(fmt)

    ph_spans = []
    for k, v in reps.items():
        for m in re.finditer(re.escape(k), full):
            ph_spans.append((m.start(), m.end(), k, str(v)))
    ph_spans.sort()

    segments = []
    pos = 0
    for ph_start, ph_end, key, value in ph_spans:
        for i, ch in enumerate(full[pos:ph_start]):
            fmt = char_fmts[pos + i] if (pos + i) < len(char_fmts) else (None,) * 5
            if segments and segments[-1][1] == fmt:
                segments[-1] = (segments[-1][0] + ch, fmt)
            else:
                segments.append((ch, fmt))
        ph_fmt = char_fmts[ph_start] if ph_start < len(char_fmts) else (None,) * 5
        if segments and segments[-1][1] == ph_fmt:
            segments[-1] = (segments[-1][0] + value, ph_fmt)
        else:
            segments.append((value, ph_fmt))
        pos = ph_end

    for i, ch in enumerate(full[pos:]):
        fmt = char_fmts[pos + i] if (pos + i) < len(char_fmts) else (None,) * 5
        if segments and segments[-1][1] == fmt:
            segments[-1] = (segments[-1][0] + ch, fmt)
        else:
            segments.append((ch, fmt))

    if not segments:
        return

    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    p_elem = para._element
    for r_elem in p_elem.findall(f"{{{W_NS}}}r"):
        p_elem.remove(r_elem)

    pPr = p_elem.find(f"{{{W_NS}}}pPr")
    insert_after = pPr

    for text, (bold, underline, italic, font_size, font_name) in segments:
        if not text:
            continue
        r_elem = _make_run_elem(text, bold, underline, italic, font_size, font_name)
        if insert_after is not None:
            insert_after.addnext(r_elem)
        else:
            p_elem.insert(0, r_elem)
        insert_after = r_elem


def fill_template(template_bytes: bytes, reps: dict) -> bytes:
    doc = Document(io.BytesIO(template_bytes))

    def proc(paras):
        for p in paras:
            replace_in_para(p, reps)

    proc(doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                proc(cell.paragraphs)
    for sec in doc.sections:
        if sec.header:
            proc(sec.header.paragraphs)
            for tbl in sec.header.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        proc(cell.paragraphs)
        if sec.footer:
            proc(sec.footer.paragraphs)
            for tbl in sec.footer.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        proc(cell.paragraphs)

    mid = io.BytesIO()
    doc.save(mid)
    mid_bytes = mid.getvalue()

    with zipfile.ZipFile(io.BytesIO(mid_bytes), "r") as zin:
        names = zin.namelist()
        files = {name: zin.read(name) for name in names}

    xml = files["word/document.xml"].decode("utf-8")

    for k, v in reps.items():
        v_xml = (
            str(v)
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
        )
        xml = xml.replace(k, v_xml)

    files["word/document.xml"] = xml.encode("utf-8")

    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)
    return out.getvalue()


# ─── SESSION STATE ────────────────────────────────────────────────────────────
for key, val in [
    ("saved_documents", []),
    ("template_bytes", None),
    ("df", None),
    ("df_source", ""),
]:
    if key not in st.session_state:
        st.session_state[key] = val

# ─── HEADER ───────────────────────────────────────────────────────────────────
st.markdown("""
<div class="main-header">
  <h1>📄 Generator Surat Penagihan</h1>
  <p>KJPP Suwendho Rinaldy dan Rekan &nbsp;·&nbsp; Kantor Cabang Jakarta Kalibata</p>
</div>
""", unsafe_allow_html=True)

# ─── SIDEBAR ──────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("### ⚙️ Setup Aplikasi")

    st.markdown("**📎 Template Surat (.docx)**")
    uploaded_tpl = st.file_uploader(
        "Upload template", type=["docx"], label_visibility="collapsed"
    )
    if uploaded_tpl:
        st.session_state.template_bytes = uploaded_tpl.read()
        st.success(f"✅ {uploaded_tpl.name}")
    elif st.session_state.template_bytes:
        st.info("📄 Template tersimpan dari upload sebelumnya")
    else:
        st.warning("⚠️ Upload template .docx dulu")

    st.divider()

    st.markdown("**📊 Sumber Data Pemberi Tugas**")
    data_source = st.radio(
        "Pilih sumber",
        ["Upload File (CSV/Excel)", "Google Sheets URL (Public)"],
        label_visibility="collapsed",
    )

    if data_source == "Upload File (CSV/Excel)":
        uploaded_data = st.file_uploader(
            "Upload CSV / Excel",
            type=["csv", "xlsx", "xls"],
            label_visibility="collapsed",
        )
        if uploaded_data:
            try:
                if uploaded_data.name.lower().endswith(".csv"):
                    raw_df = pd.read_csv(uploaded_data)
                else:
                    # engine='openpyxl' — eksplisit untuk Streamlit Cloud
                    raw_df = pd.read_excel(uploaded_data, engine="openpyxl")
                df = normalize_df(raw_df)
                st.session_state.df = df
                st.session_state.df_source = uploaded_data.name
                st.success(f"✅ {len(df):,} baris dari {uploaded_data.name}")
            except Exception as e:
                st.error(f"❌ Gagal memuat file: {e}")
    else:
        st.caption(
            "Pastikan spreadsheet sudah di-share sebagai "
            "'Anyone with the link can view'"
        )
        gs_url = st.text_input(
            "Paste URL Google Sheets",
            value=f"https://docs.google.com/spreadsheets/d/{SPREADSHEET_ID}/",
            label_visibility="collapsed",
        )
        if st.button("🔄 Muat Data", use_container_width=True):
            with st.spinner("Mengambil data dari Google Sheets…"):
                try:
                    raw_df = load_from_csv_url(gs_url)
                    df = normalize_df(raw_df)
                    st.session_state.df = df
                    st.session_state.df_source = "Google Sheets"
                    st.success(f"✅ {len(df):,} baris dimuat")
                except requests.exceptions.HTTPError as e:
                    status = e.response.status_code if e.response else "?"
                    st.error(f"❌ HTTP {status}: Gagal mengambil sheet.")
                    if status in (401, 403):
                        st.info(
                            "💡 Sheet belum di-share publik. "
                            "Buka Google Sheets → Share → "
                            "Change to *Anyone with the link* → Viewer."
                        )
                    else:
                        st.info(
                            "💡 Coba download sheet sebagai CSV "
                            "lalu upload menggunakan opsi 'Upload File'."
                        )
                except Exception as e:
                    st.error(f"❌ Gagal: {str(e)[:300]}")
                    st.info(
                        "💡 Coba download sheet sebagai CSV "
                        "lalu upload menggunakan opsi 'Upload File'."
                    )

    if st.session_state.df is not None:
        st.caption(
            f"Source: {st.session_state.df_source} · "
            f"{len(st.session_state.df):,} baris · "
            f"{st.session_state.df['pemberi_tugas'].nunique()} klien unik"
        )

    st.divider()

    n_saved = len(st.session_state.saved_documents)
    if n_saved:
        st.markdown(f"**💾 Dokumen Tersimpan: {n_saved}**")
        if st.button("🗑️ Hapus Semua", use_container_width=True):
            st.session_state.saved_documents = []
            st.rerun()

# ─── GUARD ────────────────────────────────────────────────────────────────────
df = st.session_state.df
if df is None:
    st.info("⏳ Muat data pemberi tugas terlebih dahulu melalui sidebar.")
    with st.expander("💡 Cara Menggunakan"):
        st.markdown("""
**Langkah Setup:**
1. Upload template surat `.docx` di sidebar
2. Pilih sumber data:
   - **Upload File**: Download sheet sebagai `.xlsx` atau `.csv` lalu upload
   - **Google Sheets URL**: Pastikan sheet sudah di-set *Anyone can view*, lalu paste URL-nya
3. Isi form di bawah ini
4. Klik **Simpan Sementara** atau **Download Langsung**
5. File `.docx` yang sudah terisi placeholder siap diunduh

**Tip Google Sheets:** File → Share → Change to anyone with the link → Viewer
        """)
    st.stop()

if st.session_state.template_bytes is None:
    st.info("⏳ Upload template `.docx` di sidebar untuk melanjutkan.")
    st.stop()

# ─── MAIN FORM ────────────────────────────────────────────────────────────────
unique_pts = sorted(
    df["pemberi_tugas"].str.strip().drop_duplicates().tolist(), key=str.upper
)
col_left, col_right = st.columns([3, 2])

with col_left:
    # ── 1. Pilih Pemberi Tugas ──────────────────────────────────────────────
    with st.container():
        st.markdown(
            '<div class="section-card"><h3>1️⃣ Pilih Pemberi Tugas</h3>',
            unsafe_allow_html=True,
        )
        selected_pt = st.selectbox(
            "Pemberi Tugas", ["— Pilih —"] + unique_pts, key="sel_pt"
        )
        if selected_pt == "— Pilih —":
            st.markdown("</div>", unsafe_allow_html=True)
            with col_right:
                st.info("👈 Pilih pemberi tugas terlebih dahulu")
            st.stop()

        df_pt = df[df["pemberi_tugas"].str.strip() == selected_pt.strip()].copy()

        # Pemilihan bertingkat untuk No. Proposal & Nama File
        no_proposals = (
            df_pt["no_proposal"]
            .astype(str)
            .replace(r"(?i)nan", "(kosong)", regex=True)
            .fillna("(kosong)")
            .unique()
            .tolist()
        )
        sel_no = st.selectbox("Pilih No. Proposal", no_proposals, key="sel_proposal")

        df_pt = df_pt[
            df_pt["no_proposal"]
            .astype(str)
            .replace(r"(?i)nan", "(kosong)", regex=True)
            .fillna("(kosong)")
            == sel_no
        ]

        sel_file = ""
        if "nama_file" in df_pt.columns:
            nama_files = (
                df_pt["nama_file"]
                .astype(str)
                .replace(r"(?i)nan", "(kosong)", regex=True)
                .fillna("(kosong)")
                .unique()
                .tolist()
            )
            sel_file = st.selectbox(
                "Pilih Nama File", nama_files, key="sel_nama_file"
            )
            df_pt = df_pt[
                df_pt["nama_file"]
                .astype(str)
                .replace(r"(?i)nan", "(kosong)", regex=True)
                .fillna("(kosong)")
                == sel_file
            ]

        row = df_pt.iloc[0]
        st.markdown("</div>", unsafe_allow_html=True)

    # ── 2. Data dari Spreadsheet ────────────────────────────────────────────
    with st.container():
        st.markdown(
            '<div class="section-card"><h3>2️⃣ Data dari Spreadsheet</h3>',
            unsafe_allow_html=True,
        )

        def sv(field):
            """Ambil nilai dari row, kembalikan string kosong jika NaN."""
            val = row.get(field, "")
            if pd.isna(val) or str(val).strip().lower() == "nan":
                return ""
            val_str = str(val).strip()
            # Bersihkan kode_pos yang berakhiran ".0" karena konversi pandas
            if field == "kode_pos" and val_str.endswith(".0"):
                return val_str[:-2]
            return val_str

        st.caption("✏️ Data di bawah ini dapat diedit manual jika diperlukan.")

        # Key unik per baris terpilih agar nilai default ter-reset
        # saat pemberi tugas / proposal / nama file berganti
        rk = f"{selected_pt}|{sel_no}|{sel_file}"

        c1, c2 = st.columns(2)
        with c1:
            ed_no_proposal = st.text_input(
                "No. Proposal", value=sv("no_proposal"), key=f"ed_no_proposal_{rk}"
            )
            ed_tanggal_proposal = st.text_input(
                "Tanggal Proposal",
                value=sv("tanggal_proposal"),
                key=f"ed_tgl_proposal_{rk}",
            )
            ed_alamat_1 = st.text_input(
                "Alamat 1", value=sv("alamat_1"), key=f"ed_alamat_1_{rk}"
            )
            ed_kota = st.text_input("Kota", value=sv("kota"), key=f"ed_kota_{rk}")
        with c2:
            ed_up = st.text_input("U.p.", value=sv("up"), key=f"ed_up_{rk}")
            ed_penugasan = st.text_input(
                "Penugasan", value=sv("penugasan"), key=f"ed_penugasan_{rk}"
            )
            ed_alamat_2 = st.text_input(
                "Alamat 2", value=sv("alamat_2"), key=f"ed_alamat_2_{rk}"
            )
            ed_kode_pos = st.text_input(
                "Kode Pos", value=sv("kode_pos"), key=f"ed_kode_pos_{rk}"
            )
        fee_default = float(row.get("proposed_fee", 0) or 0)
        ed_fee_str = st.text_input(
            "Proposed Fee (Rp)",
            value=format_rp(fee_default),
            key=f"ed_fee_{rk}",
            help="Angka saja, pemisah titik/koma diabaikan",
        )
        proposed_fee_raw = float(re.sub(r"[^\d]", "", ed_fee_str) or 0)
        st.markdown("</div>", unsafe_allow_html=True)

    # ── 3. Isian Manual ─────────────────────────────────────────────────────
    with st.container():
        st.markdown(
            '<div class="section-card"><h3>3️⃣ Isian Manual</h3>',
            unsafe_allow_html=True,
        )
        c3, c4 = st.columns(2)
        with c3:
            tgl_srt = st.date_input(
                "Tgl_Srt (Tanggal Surat)", value=date.today(), key="tgl_srt"
            )
            auto_kode_pt = extract_kode_pt(ed_no_proposal)
            kode_pt = st.text_input(
                "Kode_PT",
                value=auto_kode_pt,
                placeholder="Contoh: BBRI, MS, UNVR…",
                key=f"kode_pt_{ed_no_proposal}",
                help=(
                    "Otomatis diambil dari No. Proposal "
                    "(segmen sebelum segmen terakhir), bisa diedit manual"
                ),
            )
            tagih_ke = st.text_input(
                "Tagih_ke",
                value="Pertama",
                placeholder="Pertama / Kedua / Ketiga",
                key="tagih_ke",
            )
        with c4:
            nomor_raw = st.number_input(
                "Nomor (3 digit)",
                min_value=1,
                max_value=999,
                value=1,
                step=1,
                key="nomor_raw",
            )
            persentase = st.number_input(
                "Persentase (%)",
                min_value=0.0,
                max_value=100.0,
                value=100.0,
                step=5.0,
                format="%.0f",
                key="persentase",
            )
            title_up = st.text_input(
                "title_Up",
                value="Bapak/Ibu",
                placeholder="Bapak / Ibu",
                key="title_up",
            )
        bank_sel = st.selectbox(
            "Bank dan Nomor Rekening", list(BANK_OPTIONS.keys()), key="bank_sel"
        )
        st.markdown("</div>", unsafe_allow_html=True)

    # ── 4. Kalkulasi Otomatis ───────────────────────────────────────────────
    with st.container():
        st.markdown(
            '<div class="section-card"><h3>4️⃣ Kalkulasi Otomatis</h3>',
            unsafe_allow_html=True,
        )
        pct = persentase / 100.0
        fee_tagih = proposed_fee_raw * pct
        dpp = fee_tagih * (11 / 12)
        ppn = dpp * 0.12
        jumlah = fee_tagih + ppn
        nomor_str = f"{int(nomor_raw):03d}"
        tgl_yymmdd = tgl_srt.strftime("%y%m%d")
        nomor_srt = f"{tgl_yymmdd}.{nomor_str}"
        j_terbilang = terbilang_rupiah(jumlah)

        calc_items = [
            ("Nomor_Srt", nomor_srt),
            ("Fee_Tagih", f"Rp {format_rp(fee_tagih)}"),
            ("DPP (Fee × 11/12)", f"Rp {format_rp(dpp)}"),
            ("PPN (12% × DPP)", f"Rp {format_rp(ppn)}"),
            ("Jumlah (Fee + PPN)", f"Rp {format_rp(jumlah)}"),
            ("Terbilang", j_terbilang),
        ]
        for label, val in calc_items:
            a, b = st.columns([1, 2])
            a.caption(label)
            b.markdown(f"**{val}**")
        st.markdown("</div>", unsafe_allow_html=True)

    # ── 5. Buttons ──────────────────────────────────────────────────────────
    pct_label = (
        f"{int(persentase)}%"
        if persentase == int(persentase)
        else f"{persentase:.1f}%"
    )
    bank_info = BANK_OPTIONS[bank_sel]

    def build_reps():
        penugasan_asli = ed_penugasan
        reps = {
            "{{Nomor_Srt}}": nomor_srt,
            "{{Kode_PT}}": kode_pt,
            "{{Tgl_Srt}}": format_tgl(tgl_srt),
            "{{pemberi_tugas}}": sv("pemberi_tugas") or selected_pt,
            "{{alamat_1}}": ed_alamat_1,
            "{{alamat_2}}": ed_alamat_2,
            "{{kota}}": ed_kota,
            "{{kode_pos}}": ed_kode_pos,
            "{{up}}": ed_up,
            # Variasi tagih_ke
            "{{tagih_ke}}": str(tagih_ke).lower(),
            "{{Tagih_ke}}": str(tagih_ke).title(),
            "{{TAGIH_KE}}": str(tagih_ke).upper(),
            # Variasi penugasan
            "{{penugasan}}": penugasan_asli.lower(),
            "{{Penugasan}}": penugasan_asli.title(),
            "{{PENUGASAN}}": penugasan_asli.upper(),
            "{{no_proposal}}": ed_no_proposal,
            "{{tanggal_proposal}}": ed_tanggal_proposal,
            "{{proposed_fee}}": format_rp(proposed_fee_raw),
            "{{persentase}}": pct_label,
            "{{Fee_Tagih}}": format_rp(fee_tagih),
            "{{DPP}}": format_rp(dpp),
            "{{PPN}}": format_rp(ppn),
            "{{Jumlah}}": format_rp(jumlah),
            "{{Jumlah_Terbilang}}": j_terbilang,
            "{{Bank}}": bank_info["bank"],
            "{{Norek}}": bank_info["norek"],
            "{{title_Up}}": title_up,
            "{{title_up}}": title_up,
        }
        if "nama_file" in df_pt.columns:
            reps["{{nama_file}}"] = sv("nama_file")
        return reps

    def make_fname():
        pt_clean = re.sub(r"[^\w\s-]", "", selected_pt).strip().replace(" ", "_")[:35]
        return f"{nomor_srt}-SK-OR-{kode_pt or 'XX'}-{pt_clean}-{tagih_ke}.docx"

    col_b1, col_b2 = st.columns(2)
    with col_b1:
        if st.button(
            "💾 Simpan Sementara", use_container_width=True, type="primary"
        ):
            if not kode_pt.strip():
                st.error("Isi Kode_PT terlebih dahulu!")
            else:
                try:
                    filled = fill_template(
                        st.session_state.template_bytes, build_reps()
                    )
                    fname = make_fname()
                    st.session_state.saved_documents.append(
                        {
                            "filename": fname,
                            "bytes": filled,
                            "meta": {
                                "pt": selected_pt,
                                "nomor_srt": nomor_srt,
                                "jumlah": f"Rp {format_rp(jumlah)}",
                                "tagih_ke": tagih_ke,
                                "at": datetime.now().strftime("%H:%M:%S"),
                            },
                        }
                    )
                    st.success(f"✅ Tersimpan: **{fname}**")
                except Exception as e:
                    st.error(f"Gagal menyimpan: {e}")

    with col_b2:
        if kode_pt.strip():
            try:
                single = fill_template(st.session_state.template_bytes, build_reps())
                st.download_button(
                    "⬇️ Download Langsung",
                    data=single,
                    file_name=make_fname(),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception:
                st.button(
                    "⬇️ Download Langsung", disabled=True, use_container_width=True
                )
        else:
            st.button(
                "⬇️ Download Langsung",
                disabled=True,
                use_container_width=True,
                help="Isi Kode_PT dulu",
            )

# ─── RIGHT COLUMN: Preview + Saved Docs ──────────────────────────────────────
with col_right:
    st.markdown("### 📋 Preview Surat")
    bank_info_p = BANK_OPTIONS[bank_sel]
    pct_label_p = (
        f"{int(persentase)}%"
        if persentase == int(persentase)
        else f"{persentase:.1f}%"
    )
    st.markdown(
        f"""
<div class="preview-box">
<b>No.:</b> {nomor_srt}/SRR-JK/SK-OR/{kode_pt or '…'}<br>
<b>Tgl:</b> {format_tgl(tgl_srt)}<br><br>
<b>Kepada Yth.</b><br>
{selected_pt}<br>
{ed_alamat_1}<br>
{ed_alamat_2}<br>
{ed_kota} {ed_kode_pos}<br>
<small>U.p.: {ed_up}</small><br><br>
<b>Hal: Penagihan {tagih_ke}</b><br>
{ed_penugasan}<br><br>
<table width="100%" style="font-size:0.85rem;border-collapse:collapse">
<tr><td>Fee ({pct_label_p} × {format_rp(proposed_fee_raw)})</td>
    <td align="right"><b>Rp {format_rp(fee_tagih)}</b></td></tr>
<tr><td>DPP (× 11/12)</td>
    <td align="right">Rp {format_rp(dpp)}</td></tr>
<tr><td>PPN 12%</td>
    <td align="right">Rp {format_rp(ppn)}</td></tr>
<tr style="border-top:2px solid #2e86c1;font-weight:bold">
    <td>Total</td>
    <td align="right">Rp {format_rp(jumlah)}</td></tr>
</table>
<br><small><i>( {j_terbilang} )</i></small><br><br>
<small>Rekening: {bank_info_p['bank']} No. {bank_info_p['norek']}</small>
</div>
""",
        unsafe_allow_html=True,
    )

    st.markdown("---")

    st.markdown("### 📥 Dokumen Tersimpan")
    docs = st.session_state.saved_documents
    if not docs:
        st.info("Belum ada. Tekan **Simpan Sementara** untuk menambahkan.")
    else:
        for i, d in enumerate(docs):
            m = d["meta"]
            with st.expander(
                f"**{i+1}. {m['pt'][:30]}** — {m['nomor_srt']}", expanded=False
            ):
                st.caption(
                    f"Tagih ke: {m['tagih_ke']} | "
                    f"Jumlah: {m['jumlah']} | "
                    f"Simpan: {m['at']}"
                )
                c_dl, c_rm = st.columns([3, 1])
                with c_dl:
                    st.download_button(
                        "⬇️ Download",
                        data=d["bytes"],
                        file_name=d["filename"],
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"dl_{i}",
                        use_container_width=True,
                    )
                with c_rm:
                    if st.button("🗑️", key=f"rm_{i}", help="Hapus"):
                        st.session_state.saved_documents.pop(i)
                        st.rerun()

        if len(docs) > 1:
            st.markdown("---")
            zip_buf = io.BytesIO()
            with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
                for d in docs:
                    zf.writestr(d["filename"], d["bytes"])
            ts = datetime.now().strftime("%Y%m%d_%H%M")
            st.download_button(
                f"📦 Download SEMUA ({len(docs)} file) → ZIP",
                data=zip_buf.getvalue(),
                file_name=f"SuratPenagihan_{ts}.zip",
                mime="application/zip",
                use_container_width=True,
                type="primary",
            )

# ─── FOOTER ───────────────────────────────────────────────────────────────────
st.markdown("---")
st.caption(
    "KJPP Suwendho Rinaldy dan Rekan · Kalibata · "
    "Generator Surat Penagihan v1.1"
)
